"""Generate two Word versions of the demo script."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

PURPLE = RGBColor(0x2d, 0x1b, 0x69)
VIOLET = RGBColor(0x7c, 0x3a, 0xed)
AMBER = RGBColor(0x92, 0x40, 0x00)
AMBER_BG = RGBColor(0xff, 0xf8, 0xe1)
WHITE = RGBColor(0xff, 0xff, 0xff)
GREY = RGBColor(0x6b, 0x72, 0x80)
RED = RGBColor(0xdc, 0x26, 0x26)

BEATS = [
    {
        "num": "Beat 1",
        "title": "Open",
        "time": "2 min",
        "where": "/dashboard — nothing clicked yet",
        "cut": "No",
        "actions": [
            "Nothing clicked — stay on /dashboard",
        ],
        "key_line": '"Thanks for making time again, Sarah. We\'ve taken what you shared — alongside what other counsellors and students told us — and built a platform around exactly those moments. Two sides: a student platform and a counsellor platform. About fifteen minutes, then we open it up. Jump in any time."',
    },
    {
        "num": "Beat 2a",
        "title": "Profile",
        "time": "45 sec",
        "where": "Click Profile in top nav",
        "cut": "No",
        "actions": [
            "Click Profile (top nav)",
            "Show: Pathway status pill — 'Pathways open'",
        ],
        "key_line": '"Roughly one in five students quietly block themselves from degree options through subject choices made years earlier — this catches it early."',
    },
    {
        "num": "Beat 2b",
        "title": "Explore + Course Detail",
        "time": "2 min",
        "where": "Click Explore",
        "cut": "No",
        "actions": [
            "Click Explore (top nav)",
            "Show: Reach / Match / Safe cards with fit scores",
            "Click Cambridge Computer Science (Reach, top of page)",
            "Scroll to campus and city life section",
        ],
        "key_line": '"One consolidated page replaces thirty browser tabs. A lot of our students are based in India, Southeast Asia, Nigeria — they can see international student %, city life, cost of living, career prospects, all in one place."',
    },
    {
        "num": "Beat 2c",
        "title": "Toolbox — Essay Workshop LIVE",
        "time": "2.5 min",
        "where": "Click Toolbox → Essay Workshop",
        "cut": "Can trim to 1.5 min if tight",
        "actions": [
            "Click Toolbox → Essay Workshop",
            'Clear textarea. Type: "Science has always been my way of making sense of the world."',
            "Click Get Feedback",
            "Show: structured feedback panel",
        ],
        "key_line": '"Immediate, structured feedback — not writing it for them, guiding them to write something great."',
    },
    {
        "num": "Beat 2d",
        "title": "Applications",
        "time": "1 min",
        "where": "Click Applications",
        "cut": "No",
        "actions": [
            "Click Applications (top nav)",
            "Show: 3 urgent items at top (What's next)",
            "Show: full list with status, tier, progress",
        ],
        "key_line": '"The page leads with the three most urgent things across all applications — they always know what to do today."',
    },
    {
        "num": "Beat 3",
        "title": "The Segue",
        "time": "1 min",
        "where": "Pivot moment — practise this click",
        "cut": "No",
        "actions": [
            "Click violet Need help pill (Cambridge card, What's next list)",
            "Show: modal pre-filled, 'AI draft · edit before sending' label",
            "Click Send to counsellor",
            'Toast: "Sent — Sarah will respond shortly"',
            "Click Faculty view pill (top-right, violet)",
        ],
        "key_line": '"The platform knows what they\'re working on, what stage they\'re at, pre-fills a structured request. Let me flip over to your side."',
    },
    {
        "num": "Beat 4",
        "title": "Visibility Unlock",
        "time": "2.5 min",
        "where": "/counsellor — pause a beat before clicking",
        "cut": "No",
        "actions": [
            "DON'T CLICK — point out: bell with red 1 badge",
            "Point out: Help requests widget, Greg's request at top",
            "Click bell → click Greg's request row",
            "Drawer slides in from right",
            "Click Notes tab → type note → Save note",
            "Click Meeting tab → Propose",
            "Click Thread tab → type reply → Send",
            "Click X to close drawer",
            "Pause on at-risk panel",
        ],
        "key_line": '"The students who genuinely need help often aren\'t the ones asking. They go invisible until it\'s too late. This surfaces them automatically. It\'s not about replacing your judgement — it\'s about making sure your judgement gets applied to the right students."',
    },
    {
        "num": "Beat 5",
        "title": "Documents",
        "time": "2 min",
        "where": "Counsellor sub-nav → Documents",
        "cut": "Can cut short — Sarah will get it fast",
        "actions": [
            "Click Documents (counsellor sub-nav)",
            "Click Overdue filter pill",
            "Click Nudge teacher on Mohammed Al-Rashid's A-Level Transcript",
            'Show: toast, status flips to "Nudge sent"',
        ],
        "key_line": '"Twenty separate email threads, a spreadsheet to track what\'s outstanding — you nudge the right person through the platform. Everything\'s logged, so you can see at a glance what\'s blocked."',
    },
    {
        "num": "Beat 6",
        "title": "Student Detail — THE HEADLINE",
        "time": "3 min",
        "where": "Students → Aarav Sharma",
        "cut": "⚠️  NEVER CUT — this is the differentiator",
        "actions": [
            "Click Students (counsellor sub-nav) → Aarav Sharma",
            "Show header: name, school, grades, 2×2 stats, completion bar",
            "Click Notes tab — show seed entries",
            "Click Timeline tab — SLOW DOWN HERE",
        ],
        "key_line": '"Counsellors told us they usually only see the snapshot. So when you sit down with them, you don\'t have to start from scratch — you can see the journey. I genuinely think this is the thing other tools don\'t have. UCAS, Cialfo — they\'re snapshot tools. This is the only place you see why a student\'s choices look the way they do."',
    },
    {
        "num": "Beat 7",
        "title": "Analytics",
        "time": "1.5 min",
        "where": "Counsellor sub-nav → Analytics",
        "cut": "Cut first if running over",
        "actions": [
            "Click Analytics (counsellor sub-nav)",
            "Show widget grid",
            "Click Download Report → Cancel print dialog",
        ],
        "key_line": '"Year-on-year comparisons. Downloadable. So when leadership asks how this year\'s cohort did, you have something other than memory and a spreadsheet."',
    },
    {
        "num": "Beat 8",
        "title": "Integrations Teaser",
        "time": "30 sec",
        "where": "No clicking — voiceover only",
        "cut": "Can fold into Beat 9",
        "actions": [
            "No clicking",
        ],
        "key_line": '"We\'re building this to sit alongside the tools you already use. Gmail, calendar, document storage. You don\'t have to leave the platform, but the platform doesn\'t try to replace everything either."',
    },
    {
        "num": "Beat 9",
        "title": "Close",
        "time": "1.5 min",
        "where": "",
        "cut": "No",
        "actions": [
            "Stop talking after closing line",
            "Let Sarah go first",
        ],
        "key_line": '"That\'s the demo. We\'re really trying to understand whether what we\'ve built actually fits how you work. The most valuable thing now is your honest reaction — what\'s resonating, what\'s not, and what\'s missing."',
    },
]

PREP = {
    "url": "https://ascenda-ashy.vercel.app",
    "login": "greg@workiflow.com / AscendaDemo!2026",
    "browser": "Chrome, clean profile, full screen, 100% zoom",
    "backup": "Second tab pre-logged in",
}

QA = [
    ("Application status updates?", "Students update it themselves for MVP. Email-sync is in flight — auto-detects updates from university inboxes."),
    ("Is the AI assistant real?", "Structured template today, fills from student context. Full LLM version in flight — same input, richer language."),
    ("Gmail / Calendar depth?", "Real OAuth on the roadmap. What you saw is the connection point — when live, meeting proposals go into your real calendar."),
    ("Reach / Match / Safe accuracy?", "Profile vs entry requirements, historical admissions data, university characteristics. Fully decomposable — they see the breakdown."),
    ("UCAS integration?", "UCAS has no API. Student runs UCAS in parallel, logs outcomes here. Email-sync longer term."),
    ("Parent access?", "Not in this build — on the roadmap. School controls who sees what."),
]

BREAKS = [
    ("Faculty view pill missing", "Hard refresh once"),
    ("Bell stays at 0 after Send", "Wait 5 sec — realtime fires in 1s, 4s poll catches it"),
    ("Cambridge page 404", "Fall back to Imperial Computing or UCL CS"),
    ("Essay workshop returns nothing", 'Type a longer sentence, retry. If still broken: "let me move on"'),
    ("Send does nothing", 'Console → if "row level security" errors, fall back to screenshots'),
    ("Matches page slow", "Let it load — don't open extra tabs (recalculates cache)"),
    ("Production URL down", "Switch to backup tab"),
]


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_run(para, text, bold=False, italic=False, size=None, color=None):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def set_para_spacing(para, before=0, after=0, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = Pt(line)


# ─────────────────────────────────────────────
# VERSION 1: One beat per page
# ─────────────────────────────────────────────

def build_v1():
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # ── Cover page ──
    p = doc.add_paragraph()
    set_para_spacing(p, before=40, after=4)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, "ASCENDA", bold=True, size=32, color=PURPLE)

    p = doc.add_paragraph()
    set_para_spacing(p, after=2)
    add_run(p, "Demo Script  ·  Sunday 18 May 2026  ·  Sarah", size=13, color=GREY)

    p = doc.add_paragraph()
    set_para_spacing(p, after=20)
    add_run(p, "Target: 18 min demo · 12 min Q&A", bold=True, size=11, color=VIOLET)

    # Prep table
    p = doc.add_paragraph()
    set_para_spacing(p, before=8, after=6)
    add_run(p, "BEFORE YOU START", bold=True, size=9, color=GREY)

    tbl = doc.add_table(rows=4, cols=2)
    tbl.style = 'Table Grid'
    rows_data = [
        ("URL", PREP["url"]),
        ("Login", PREP["login"]),
        ("Browser", PREP["browser"]),
        ("Backup", PREP["backup"]),
    ]
    for i, (label, val) in enumerate(rows_data):
        row = tbl.rows[i]
        row.cells[0].width = Cm(2.8)
        row.cells[1].width = Cm(13)
        set_cell_bg(row.cells[0], "EDE9FE")
        cp = row.cells[0].paragraphs[0]
        add_run(cp, label, bold=True, size=9, color=PURPLE)
        vp = row.cells[1].paragraphs[0]
        add_run(vp, val, size=9)

    # Flow
    p = doc.add_paragraph()
    set_para_spacing(p, before=16, after=4)
    add_run(p, "FLOW  ", bold=True, size=9, color=GREY)
    add_run(p, "Open → Profile → Explore (campus life) → Toolbox (essay LIVE) → Applications → Need Help → Faculty view → Bell → Drawer → At-risk → Documents → Student detail → Analytics → Integrations → Close", size=9, color=GREY)

    doc.add_page_break()

    # ── Beat pages ──
    for beat in BEATS:
        never_cut = "NEVER CUT" in beat["cut"]

        # Header bar via table (full width, purple bg)
        tbl = doc.add_table(rows=1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        tbl.columns[0].width = Cm(11)
        tbl.columns[1].width = Cm(4.8)
        set_cell_bg(tbl.cell(0, 0), "2d1b69")
        set_cell_bg(tbl.cell(0, 1), "2d1b69")

        lp = tbl.cell(0, 0).paragraphs[0]
        lp.paragraph_format.space_before = Pt(6)
        lp.paragraph_format.space_after = Pt(6)
        add_run(lp, f"{beat['num']}  ·  {beat['title']}", bold=True, size=15, color=WHITE)

        rp = tbl.cell(0, 1).paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rp.paragraph_format.space_before = Pt(6)
        rp.paragraph_format.space_after = Pt(6)
        add_run(rp, beat["time"], bold=False, size=12, color=RGBColor(0xc4, 0xb5, 0xfd))

        # Where / cut line
        p = doc.add_paragraph()
        set_para_spacing(p, before=5, after=10)
        if beat["where"]:
            add_run(p, beat["where"], italic=True, size=9.5, color=GREY)
        if never_cut:
            add_run(p, "   ⚠ NEVER CUT", bold=True, size=9.5, color=RED)
        elif beat["cut"] != "No":
            add_run(p, f"   Cut: {beat['cut']}", size=9, color=GREY)

        # Actions
        p = doc.add_paragraph()
        set_para_spacing(p, before=2, after=6)
        add_run(p, "ACTIONS", bold=True, size=8.5, color=VIOLET)

        for action in beat["actions"]:
            bullet = doc.add_paragraph(style='List Bullet')
            set_para_spacing(bullet, before=2, after=2)
            bullet.paragraph_format.left_indent = Cm(0.4)
            is_show = action.lower().startswith("show")
            is_warn = action.startswith("DON'T") or action.startswith("⚠") or action.startswith("SLOW") or action.startswith("Pause")
            color = GREY if is_show else (RED if is_warn else None)
            run = bullet.add_run(action)
            run.bold = not is_show
            run.font.size = Pt(10.5)
            if color:
                run.font.color.rgb = color

        # Spoken line box via 1-cell table
        p = doc.add_paragraph()
        set_para_spacing(p, before=14, after=4)
        add_run(p, "SAY", bold=True, size=8.5, color=VIOLET)

        qt = doc.add_table(rows=1, cols=1)
        set_cell_bg(qt.cell(0, 0), "F5F3FF")
        qp = qt.cell(0, 0).paragraphs[0]
        qp.paragraph_format.space_before = Pt(8)
        qp.paragraph_format.space_after = Pt(8)
        qp.paragraph_format.left_indent = Cm(0.3)
        qp.paragraph_format.right_indent = Cm(0.3)
        add_run(qp, beat["key_line"], italic=True, size=10.5, color=PURPLE)

        doc.add_page_break()

    # ── Q&A page ──
    p = doc.add_paragraph()
    set_para_spacing(p, before=0, after=8)
    add_run(p, "LIKELY QUESTIONS", bold=True, size=14, color=PURPLE)

    tbl = doc.add_table(rows=len(QA) + 1, cols=2)
    tbl.style = 'Table Grid'
    set_cell_bg(tbl.cell(0, 0), "EDE9FE")
    set_cell_bg(tbl.cell(0, 1), "EDE9FE")
    add_run(tbl.cell(0, 0).paragraphs[0], "Question", bold=True, size=9, color=PURPLE)
    add_run(tbl.cell(0, 1).paragraphs[0], "Answer", bold=True, size=9, color=PURPLE)
    for i, (q, a) in enumerate(QA, 1):
        add_run(tbl.cell(i, 0).paragraphs[0], q, bold=True, size=9.5)
        add_run(tbl.cell(i, 1).paragraphs[0], a, size=9.5)

    doc.add_paragraph()

    # ── Breaks page ──
    p = doc.add_paragraph()
    set_para_spacing(p, before=16, after=8)
    add_run(p, "IF SOMETHING BREAKS", bold=True, size=14, color=PURPLE)

    tbl = doc.add_table(rows=len(BREAKS) + 1, cols=2)
    tbl.style = 'Table Grid'
    set_cell_bg(tbl.cell(0, 0), "FEF2F2")
    set_cell_bg(tbl.cell(0, 1), "FEF2F2")
    add_run(tbl.cell(0, 0).paragraphs[0], "Problem", bold=True, size=9, color=RED)
    add_run(tbl.cell(0, 1).paragraphs[0], "Fix", bold=True, size=9, color=RED)
    for i, (prob, fix) in enumerate(BREAKS, 1):
        add_run(tbl.cell(i, 0).paragraphs[0], prob, bold=True, size=9.5)
        add_run(tbl.cell(i, 1).paragraphs[0], fix, size=9.5)

    # ── Timing ──
    doc.add_paragraph()
    p = doc.add_paragraph()
    set_para_spacing(p, before=16, after=8)
    add_run(p, "TIMING GUIDE", bold=True, size=14, color=PURPLE)

    timing = [
        ("1 · Open", "2 min", "No"),
        ("2 · Student platform", "6 min", "Can trim essay workshop to 1.5 min"),
        ("3 · Segue", "1 min", "No"),
        ("4 · Visibility (drawer)", "2.5 min", "No"),
        ("5 · Documents", "2 min", "Can cut short"),
        ("6 · Student detail", "3 min", "⚠ NEVER CUT"),
        ("7 · Analytics", "1.5 min", "Cut first if running over"),
        ("8 · Integrations", "0.5 min", "Can fold into Beat 9"),
        ("9 · Close", "1.5 min", "No"),
        ("Total", "~20 min", "Target 18, leave 12 for Q&A"),
    ]
    tbl = doc.add_table(rows=len(timing) + 1, cols=3)
    tbl.style = 'Table Grid'
    for j, hdr in enumerate(["Beat", "Time", "Cut?"]):
        set_cell_bg(tbl.cell(0, j), "EDE9FE")
        add_run(tbl.cell(0, j).paragraphs[0], hdr, bold=True, size=9, color=PURPLE)
    for i, (beat, time, cut) in enumerate(timing, 1):
        never = "NEVER" in cut
        if never:
            for j in range(3):
                set_cell_bg(tbl.cell(i, j), "FEF2F2")
        add_run(tbl.cell(i, 0).paragraphs[0], beat, bold=never, size=9.5, color=RED if never else None)
        add_run(tbl.cell(i, 1).paragraphs[0], time, bold=never, size=9.5)
        add_run(tbl.cell(i, 2).paragraphs[0], cut, bold=never, size=9.5, color=RED if never else None)

    doc.save("/Users/gregfranck/Ascenda/docs/demo-script-v1-one-per-page.docx")
    print("Saved v1")


# ─────────────────────────────────────────────
# VERSION 2: Compact two-column
# ─────────────────────────────────────────────

def build_v2():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # Header
    p = doc.add_paragraph()
    set_para_spacing(p, before=0, after=2)
    add_run(p, "ASCENDA  ·  Demo Script  ·  18 May 2026  ·  Sarah  ·  Target 18 min demo + 12 min Q&A", bold=True, size=10, color=PURPLE)

    # Prep row
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Table Grid'
    prep_items = [("URL", PREP["url"]), ("Login", PREP["login"]), ("Browser", PREP["browser"]), ("Backup", PREP["backup"])]
    for i, (label, val) in enumerate(prep_items):
        set_cell_bg(tbl.cell(0, i), "EDE9FE")
        p2 = tbl.cell(0, i).paragraphs[0]
        add_run(p2, label + ": ", bold=True, size=8, color=PURPLE)
        add_run(p2, val, size=8)

    p = doc.add_paragraph()
    set_para_spacing(p, before=4, after=8)
    add_run(p, "Flow: Open → Profile → Explore (campus life) → Toolbox (essay LIVE) → Applications → Need Help → Faculty view → Bell → Drawer → At-risk → Documents → Student detail → Analytics → Integrations → Close", size=8, color=GREY)

    # Beats — two-column table per beat
    for beat in BEATS:
        never_cut = "NEVER CUT" in beat["cut"]

        # Beat header (full-width merged cell)
        tbl = doc.add_table(rows=1, cols=2)
        tbl.columns[0].width = Cm(7.8)
        tbl.columns[1].width = Cm(8.0)
        set_cell_bg(tbl.cell(0, 0), "2d1b69")
        set_cell_bg(tbl.cell(0, 1), "2d1b69")

        lp = tbl.cell(0, 0).paragraphs[0]
        lp.paragraph_format.space_before = Pt(5)
        lp.paragraph_format.space_after = Pt(5)
        add_run(lp, f"{beat['num']}  ·  {beat['title']}", bold=True, size=11.5, color=WHITE)

        rp = tbl.cell(0, 1).paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rp.paragraph_format.space_before = Pt(5)
        rp.paragraph_format.space_after = Pt(5)
        cut_text = "⚠ NEVER CUT" if never_cut else (beat["cut"] if beat["cut"] != "No" else "")
        add_run(rp, beat["time"] + ("  " + cut_text if cut_text else ""), size=9.5, color=RGBColor(0xff, 0xc0, 0x40) if never_cut else RGBColor(0xc4, 0xb5, 0xfd))

        # Content row
        tbl2 = doc.add_table(rows=1, cols=2)
        tbl2.style = 'Table Grid'
        tbl2.columns[0].width = Cm(7.8)
        tbl2.columns[1].width = Cm(8.0)

        # Left: Actions
        left_cell = tbl2.cell(0, 0)
        set_cell_bg(left_cell, "FAFAF9")
        lp2 = left_cell.paragraphs[0]
        lp2.paragraph_format.space_before = Pt(5)
        add_run(lp2, "ACTIONS", bold=True, size=7.5, color=VIOLET)

        for action in beat["actions"]:
            bp = left_cell.add_paragraph()
            bp.paragraph_format.left_indent = Cm(0.2)
            bp.paragraph_format.space_before = Pt(2)
            bp.paragraph_format.space_after = Pt(1)
            is_show = action.lower().startswith("show")
            is_warn = action.startswith("DON'T") or action.startswith("⚠") or action.startswith("Pause") or action.startswith("SLOW")
            run = bp.add_run(("▸  " if not is_show else "      ↳  ") + action)
            run.font.size = Pt(9.5)
            run.bold = not is_show
            if is_show:
                run.font.color.rgb = GREY
            elif is_warn:
                run.font.color.rgb = RED

        if beat["where"]:
            wp = left_cell.add_paragraph()
            wp.paragraph_format.space_before = Pt(4)
            add_run(wp, beat["where"], italic=True, size=8, color=GREY)

        # Right: Say
        right_cell = tbl2.cell(0, 1)
        set_cell_bg(right_cell, "F5F3FF")
        rp2 = right_cell.paragraphs[0]
        rp2.paragraph_format.space_before = Pt(5)
        add_run(rp2, "SAY", bold=True, size=7.5, color=VIOLET)

        sp = right_cell.add_paragraph()
        sp.paragraph_format.space_before = Pt(4)
        sp.paragraph_format.space_after = Pt(4)
        sp.paragraph_format.left_indent = Cm(0.1)
        sp.paragraph_format.right_indent = Cm(0.2)
        add_run(sp, beat["key_line"], italic=True, size=10, color=PURPLE)

        # Spacer
        p = doc.add_paragraph()
        set_para_spacing(p, before=0, after=8)

    # Q&A
    p = doc.add_paragraph()
    set_para_spacing(p, before=8, after=6)
    add_run(p, "LIKELY QUESTIONS", bold=True, size=12, color=PURPLE)

    tbl = doc.add_table(rows=len(QA) + 1, cols=2)
    tbl.style = 'Table Grid'
    set_cell_bg(tbl.cell(0, 0), "EDE9FE")
    set_cell_bg(tbl.cell(0, 1), "EDE9FE")
    add_run(tbl.cell(0, 0).paragraphs[0], "Question", bold=True, size=8.5, color=PURPLE)
    add_run(tbl.cell(0, 1).paragraphs[0], "Answer", bold=True, size=8.5, color=PURPLE)
    for i, (q, a) in enumerate(QA, 1):
        add_run(tbl.cell(i, 0).paragraphs[0], q, bold=True, size=9)
        add_run(tbl.cell(i, 1).paragraphs[0], a, size=9)

    p = doc.add_paragraph()
    set_para_spacing(p, before=12, after=6)
    add_run(p, "IF SOMETHING BREAKS", bold=True, size=12, color=RED)

    tbl = doc.add_table(rows=len(BREAKS) + 1, cols=2)
    tbl.style = 'Table Grid'
    set_cell_bg(tbl.cell(0, 0), "FEF2F2")
    set_cell_bg(tbl.cell(0, 1), "FEF2F2")
    add_run(tbl.cell(0, 0).paragraphs[0], "Problem", bold=True, size=8.5, color=RED)
    add_run(tbl.cell(0, 1).paragraphs[0], "Fix", bold=True, size=8.5, color=RED)
    for i, (prob, fix) in enumerate(BREAKS, 1):
        add_run(tbl.cell(i, 0).paragraphs[0], prob, bold=True, size=9)
        add_run(tbl.cell(i, 1).paragraphs[0], fix, size=9)

    doc.save("/Users/gregfranck/Ascenda/docs/demo-script-v2-two-column.docx")
    print("Saved v2")


build_v1()
build_v2()
print("Done.")
