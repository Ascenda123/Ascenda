"""Generate counsellor and student presentation guides as Word docs."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PURPLE = RGBColor(0x2d, 0x1b, 0x69)
VIOLET = RGBColor(0x7c, 0x3a, 0xed)
GREEN = RGBColor(0x05, 0x5a, 0x35)
GREY = RGBColor(0x6b, 0x72, 0x80)
RED = RGBColor(0xb9, 0x1c, 0x1c)
AMBER = RGBColor(0x78, 0x35, 0x00)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def r(para, text, bold=False, italic=False, size=None, color=None):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def sp(para, before=0, after=6):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)


def section_heading(doc, text, color=PURPLE, before=18):
    p = doc.add_paragraph()
    sp(p, before=before, after=5)
    r(p, text.upper(), bold=True, size=11.5, color=color)
    # Underline via border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '7c3aed')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def pain_block(doc, pain_text, say_text, action_text=None):
    """A self-contained unit: pain point label + pain text, action, talking point."""
    # Pain label
    p = doc.add_paragraph()
    sp(p, before=10, after=2)
    r(p, "PAIN POINT  ", bold=True, size=8, color=RED)
    r(p, pain_text, size=10, italic=True, color=RED)

    # Optional action line
    if action_text:
        p2 = doc.add_paragraph()
        sp(p2, before=2, after=3)
        r(p2, "▸  ", bold=True, size=10, color=VIOLET)
        r(p2, action_text, size=10, bold=True)

    # Say box
    tbl = doc.add_table(rows=1, cols=1)
    set_cell_bg(tbl.cell(0, 0), "F5F3FF")
    qp = tbl.cell(0, 0).paragraphs[0]
    qp.paragraph_format.space_before = Pt(7)
    qp.paragraph_format.space_after = Pt(7)
    qp.paragraph_format.left_indent = Cm(0.35)
    qp.paragraph_format.right_indent = Cm(0.35)
    r(qp, "SAY:  ", bold=True, size=8.5, color=VIOLET)
    r(qp, say_text, italic=True, size=10.5, color=PURPLE)

    p3 = doc.add_paragraph()
    sp(p3, before=0, after=2)


def note_line(doc, text, color=GREY):
    p = doc.add_paragraph()
    sp(p, before=2, after=3)
    r(p, "↳  ", bold=True, size=10, color=color)
    r(p, text, size=10, italic=True, color=color)


def two_col_table(doc, rows_data, left_header, right_header, left_bg="EDE9FE", header_color=PURPLE):
    tbl = doc.add_table(rows=len(rows_data) + 1, cols=2)
    tbl.style = 'Table Grid'
    set_cell_bg(tbl.cell(0, 0), left_bg)
    set_cell_bg(tbl.cell(0, 1), left_bg)
    tbl.columns[0].width = Cm(8.2)
    tbl.columns[1].width = Cm(7.6)
    r(tbl.cell(0, 0).paragraphs[0], left_header, bold=True, size=9, color=header_color)
    r(tbl.cell(0, 1).paragraphs[0], right_header, bold=True, size=9, color=header_color)
    for i, (left, right) in enumerate(rows_data, 1):
        tbl.cell(i, 0).width = Cm(8.2)
        tbl.cell(i, 1).width = Cm(7.6)
        r(tbl.cell(i, 0).paragraphs[0], left, bold=True, size=9.5)
        r(tbl.cell(i, 1).paragraphs[0], right, size=9.5)
    p = doc.add_paragraph()
    sp(p, before=0, after=10)


# ═══════════════════════════════════════════════════════════════
# COUNSELLOR GUIDE
# ═══════════════════════════════════════════════════════════════

def build_counsellor():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # Title block
    p = doc.add_paragraph()
    sp(p, before=0, after=3)
    r(p, "Ascenda", bold=True, size=26, color=PURPLE)
    p = doc.add_paragraph()
    sp(p, before=0, after=2)
    r(p, "Counsellor & Faculty Presentation Guide", bold=True, size=13, color=VIOLET)
    p = doc.add_paragraph()
    sp(p, before=0, after=14)
    r(p, "1:1 buy-in  ·  ~20 min  ·  Feedback gathering", size=9.5, color=GREY)

    # ── PURPOSE ──
    section_heading(doc, "What this session is for", before=0)
    p = doc.add_paragraph()
    sp(p, before=6, after=4)
    r(p, "You're not demoing a finished product. You're showing a counsellor how their day could look — and then listening. Two goals: spark genuine interest, and leave with honest feedback you can act on.", size=10.5)
    p = doc.add_paragraph()
    sp(p, before=0, after=4)
    r(p, "Every feature you show should answer a problem they've already told you they have. If you don't know their problems yet, ask before you open the laptop.", size=10.5)

    # ── BEFORE THE CALL ──
    section_heading(doc, "Before the call")
    p = doc.add_paragraph()
    sp(p, before=5, after=4)
    r(p, "Know their context before you start — these shape which parts of the demo to emphasise:", size=10.5)

    context_items = [
        ("How many students do they manage?", "Determines how much the at-risk panel and bulk document nudging will resonate."),
        ("IB, A-Level, or mixed?", "Affects whether the profile → pathway pill or the grade-tracking in Timeline is more relevant."),
        ("Do they use Cialfo, UCAS Hub, or nothing?", "If Cialfo: lean into the Timeline — it's the clear gap. If nothing: lean into everything."),
        ("What part of the cycle are they in right now?", "If deadlines are imminent, lead with Documents and at-risk. If early in cycle, lead with student detail and Timeline."),
    ]
    for q, why in context_items:
        p2 = doc.add_paragraph()
        sp(p2, before=3, after=1)
        p2.paragraph_format.left_indent = Cm(0.3)
        r(p2, "▸  ", bold=True, size=10, color=VIOLET)
        r(p2, q + "  ", bold=True, size=10)
        r(p2, why, size=10, color=GREY)

    p = doc.add_paragraph()
    sp(p, before=8, after=3)
    r(p, "Technical setup:", bold=True, size=10)

    setup = [
        "Reset demo state in Supabase (SQL in the demo script) — bell must start at zero.",
        "Platform open on /dashboard, clean Chrome profile. Faculty view pill visible top-right.",
        "Second tab pre-logged-in as backup. Don't switch tabs mid-demo.",
    ]
    for s in setup:
        p2 = doc.add_paragraph()
        sp(p2, before=1, after=1)
        p2.paragraph_format.left_indent = Cm(0.3)
        r(p2, "▸  ", bold=True, size=10, color=VIOLET)
        r(p2, s, size=10)

    # ── OPENING ──
    section_heading(doc, "Opening — set the frame (2 min)")
    p = doc.add_paragraph()
    sp(p, before=6, after=4)
    r(p, "Don't open the laptop first. Open with their problem. Ask:", size=10.5)

    pain_block(
        doc,
        pain_text="Counsellors are reactive — their time gets consumed by whoever asks loudest, not whoever needs them most.",
        action_text='Ask: "What\'s the most repetitive part of the application cycle for you right now?"',
        say_text='"Before I show you anything — what\'s taking the most time and giving back the least? I want to make sure what I show you is actually relevant to your situation."',
    )
    note_line(doc, "Let them answer in full. Their answer is your anchor. When you show a feature, connect it back to exactly what they just said.", color=GREY)

    p = doc.add_paragraph()
    sp(p, before=8, after=3)
    r(p, "Then frame the session:", size=10.5)

    tbl = doc.add_table(rows=1, cols=1)
    set_cell_bg(tbl.cell(0, 0), "F5F3FF")
    qp = tbl.cell(0, 0).paragraphs[0]
    qp.paragraph_format.space_before = Pt(7)
    qp.paragraph_format.space_after = Pt(7)
    qp.paragraph_format.left_indent = Cm(0.35)
    r(qp, "SAY:  ", bold=True, size=8.5, color=VIOLET)
    r(qp, '"Two sides — a student platform and a counsellor platform. I\'ll show the student side briefly because everything on your side flows from what students do. About fifteen minutes, then I want your honest reaction. Jump in any time."', italic=True, size=10.5, color=PURPLE)
    doc.add_paragraph()

    # ── STUDENT SIDE ──
    section_heading(doc, "Student side — show quickly (5–6 min)")
    p = doc.add_paragraph()
    sp(p, before=5, after=8)
    r(p, "The goal here is not to explain features — it's to make them feel the counsellor value on the other side. Keep moving.", size=10.5)

    pain_block(
        doc,
        pain_text="Students make subject choices years before applying that quietly close degree options — and neither the student nor the counsellor realises until it's too late.",
        action_text="Click Profile (top nav). Show the Pathway status pill.",
        say_text='"When a student fills in their profile — grades, subjects, what they want to study — the system immediately checks whether their current subject choices keep their options open. Roughly one in five students block themselves from degree options through choices made years earlier. This pill flags it before the exam window closes, while something can still be done."',
    )

    pain_block(
        doc,
        pain_text="Students waste hours across fragmented university websites and still don't have the information they need to make a confident decision.",
        action_text="Click Explore. Show Reach / Match / Safe cards. Click Cambridge CS. Scroll to campus life section.",
        say_text='"The match isn\'t arbitrary — it runs their profile against real entry requirements and university characteristics. And when they click into a course, one page replaces thirty browser tabs. The section I\'d draw your attention to: campus life, international student percentage, cost of living. A lot of your students are applying to places they\'ve never visited. This is what helps them make a genuinely informed choice."',
    )

    pain_block(
        doc,
        pain_text="Students don't know what a strong personal statement looks like — and counsellors don't have time to review every draft.",
        action_text='Click Toolbox → Essay Workshop. Type: "Science has always been my way of making sense of the world." Click Get Feedback.',
        say_text='"UCAS is 4,000 characters. Other portals are different. Students don\'t know what good looks like. This benchmarks their draft against the strongest essays from Oxford, Cambridge, Harvard — and gives them structured, actionable feedback. Not writing it for them. Giving them a standard to work toward, before it lands on your desk."',
    )

    pain_block(
        doc,
        pain_text="Students lose track of what needs doing across multiple applications on different portals with different deadlines.",
        action_text="Click Applications. Show the What's next priority board at the top.",
        say_text='"This is the student\'s central hub. The priority board at the top shows the three most urgent things across all their applications — they always know what to do today. Each application has its own task list so nothing falls through the cracks."',
    )

    p = doc.add_paragraph()
    sp(p, before=8, after=3)
    r(p, "Now pivot:", bold=True, size=10.5)
    p2 = doc.add_paragraph()
    sp(p2, before=2, after=2)
    p2.paragraph_format.left_indent = Cm(0.3)
    r(p2, "▸  ", bold=True, size=10, color=VIOLET)
    r(p2, 'Click the violet Need help pill next to Cambridge. Show the pre-filled modal. Click Send to counsellor. Then click Faculty view (top-right, violet).', size=10, bold=True)

    tbl = doc.add_table(rows=1, cols=1)
    set_cell_bg(tbl.cell(0, 0), "F5F3FF")
    qp = tbl.cell(0, 0).paragraphs[0]
    qp.paragraph_format.space_before = Pt(7)
    qp.paragraph_format.space_after = Pt(7)
    qp.paragraph_format.left_indent = Cm(0.35)
    r(qp, "SAY:  ", bold=True, size=8.5, color=VIOLET)
    r(qp, '"The platform knows what they\'re working on, what stage they\'re at — pre-fills a structured request. Rather than a blank email that tells you nothing, you get context. Let me flip over to your side."', italic=True, size=10.5, color=PURPLE)
    doc.add_paragraph()

    # ── COUNSELLOR SIDE ──
    section_heading(doc, "Counsellor side — slow down (10–12 min)")
    p = doc.add_paragraph()
    sp(p, before=5, after=8)
    r(p, "This is where the conversation should shift. Narrate less. Pause after each section and check if it's landing.", size=10.5)

    pain_block(
        doc,
        pain_text="Urgent student needs get missed because there's no single place that surfaces them — counsellors find out too late.",
        action_text="Don't click yet. Point out: bell with red 1 badge, Help requests widget with Greg's request at top. Then click bell → click row → drawer opens.",
        say_text='"This is your homepage. The notification from the student we just sent lives here — you don\'t need to check email or dig through a portal. Open the drawer and you have three things you can do without leaving the page: reply, leave a private note for yourself, or propose a meeting. Everything in one place."',
    )
    note_line(doc, "Click Notes tab. Type a note. Click Save. Click Meeting tab. Click Propose. Click Thread tab. Type a reply. Click Send. Close drawer.")

    pain_block(
        doc,
        pain_text="Counsellors spend too much time with students who ask the most, not the ones who need them most — quiet students fall through the cracks.",
        action_text="Pause on the at-risk panel below the help requests widget.",
        say_text='"This is the bit I\'d ask you to look at carefully. The pattern we hear consistently: meetings stack up, the loudest students absorb the most time, and the students who genuinely need help often aren\'t the ones asking. They go invisible until it\'s too late. This surfaces them automatically — approaching deadlines, incomplete work, applications stalled. It\'s not about replacing your judgement. It\'s about making sure your judgement gets applied to the right students."',
    )

    pain_block(
        doc,
        pain_text="Document chasing — transcripts, references, personal statements — is one of the most time-consuming and repetitive parts of the cycle.",
        action_text="Click Documents (counsellor sub-nav). Click Overdue filter pill. Click Nudge teacher on Mohammed Al-Rashid's A-Level Transcript.",
        say_text='"Document chasing normally means twenty separate email threads, a spreadsheet to track what\'s outstanding, chasing students who haven\'t replied to chase teachers who haven\'t replied. Here: overdue documents surface automatically. You nudge the right person — student, teacher, registrar — through the platform. Everything\'s logged, so you can see at a glance what\'s blocked instead of holding it all in your head."',
    )

    pain_block(
        doc,
        pain_text="1:1 time is the most valuable thing a counsellor gives a student — and the hardest to protect. Counsellors walk into meetings underprepared. When a student moves counsellors, the context is lost.",
        action_text="Click Students → click Aarav Sharma. Show header card (don't scroll). Click Notes tab. Click Timeline tab. SLOW DOWN HERE.",
        say_text='"This is what walking into a 1:1 looks like with Ascenda. Thirty seconds and you\'re fully briefed — who they are, where they\'re applying, what\'s outstanding, next deadline. The notes tab is yours: private, attached to the student, survives counsellor turnover. And the timeline — this is the thing other tools don\'t have. UCAS, Cialfo — they show you a snapshot. This shows you the journey. When a student\'s shortlist shifts after an open day, or they drop a subject, or a grade changes — it\'s all here. So when you sit down with them and ask why their list looks the way it does, you already know."',
    )

    pain_block(
        doc,
        pain_text="Leadership asks how this year's cohort compares to last year — and the answer lives in memory or a spreadsheet someone built manually.",
        action_text="Click Analytics (counsellor sub-nav). Show the widget grid. Click Download Report → Cancel print dialog.",
        say_text='"Year-on-year comparisons, built automatically as students hear back. Downloadable. If a lot of your students are applying to the same field or country, the data tells you where to run a targeted workshop. When leadership asks how this year went, you have something other than a spreadsheet."',
    )
    note_line(doc, "If running over time, cut Analytics first — it's the easiest to explain in a follow-up email.", color=AMBER)

    # ── CLOSING ──
    section_heading(doc, "Closing — hand it over (2 min)")
    p = doc.add_paragraph()
    sp(p, before=6, after=4)
    r(p, "Stop before they check out. The last thing they see should be the platform, not a slide. Then ask:", size=10.5)

    tbl = doc.add_table(rows=1, cols=1)
    set_cell_bg(tbl.cell(0, 0), "F5F3FF")
    qp = tbl.cell(0, 0).paragraphs[0]
    qp.paragraph_format.space_before = Pt(7)
    qp.paragraph_format.space_after = Pt(7)
    qp.paragraph_format.left_indent = Cm(0.35)
    r(qp, "SAY:  ", bold=True, size=8.5, color=VIOLET)
    r(qp, '"That\'s the platform. We\'re at a stage where the most valuable thing is your honest reaction — what resonates, what doesn\'t, and what\'s missing for how you actually work."', italic=True, size=10.5, color=PURPLE)
    doc.add_paragraph()

    note_line(doc, "Stop talking. Let them go first. Don't fill the silence.")

    # ── FEEDBACK QUESTIONS ──
    section_heading(doc, "Feedback questions to ask")
    p = doc.add_paragraph()
    sp(p, before=5, after=6)
    r(p, "Pick 3–4 depending on what came up. Don't ask them all — it becomes an interview. The notes column is for you to fill in during the call.", size=10)

    two_col_table(doc, [
        ("Which part would actually change how you work day-to-day?", ""),
        ("Is there anything here you'd be worried students wouldn't use?", ""),
        ("What's the biggest thing missing for your context?", ""),
        ("You mentioned [X] earlier — does what you saw change that at all?", ""),
        ("Would you want to see this before or after students have built their shortlist?", ""),
        ("What would you need to see to feel confident recommending this to a student?", ""),
        ("How does this compare to what you're using now for tracking applications?", ""),
        ("Who else at your school would need to be involved in a decision like this?", ""),
    ], "Question to ask", "Notes")

    # ── OBJECTIONS ──
    section_heading(doc, "Likely objections")
    two_col_table(doc, [
        ("We already use Cialfo / Unifrog",
         "We're not trying to replace those — they track applications. What they don't have is the longitudinal view of the student: the Timeline, the prep for 1:1s, the at-risk surfacing. That's the gap we fill."),
        ("Students won't keep it updated",
         "The student side is built so updates happen as part of normal use — the application tracker is their to-do list, not a separate admin task. Keeping it updated is using it."),
        ("How accurate is the matching?",
         "It runs against real entry requirements and university data, and it's fully decomposable — the student sees exactly why a course ranked where it did. It's not a black box."),
        ("What does it cost?",
         "We're in early access and working with schools on pricing. What's more useful right now is understanding whether this actually fits how you work — pricing follows from that conversation."),
        ("We'd need IT sign-off / data compliance",
         "It's web-based, no install required. Data is stored securely and we can walk your IT team through the architecture. Happy to send the spec."),
        ("We don't have time to onboard students onto another tool",
         "The student side is designed to be self-explanatory — profile takes ten minutes, the platform guides them from there. The counsellor side requires no student onboarding at all."),
    ], "They say…", "You say…", left_bg="FEF2F2", header_color=RED)

    # ── AFTER THE CALL ──
    section_heading(doc, "After the call — within 24 hours")
    after_items = [
        ("Send a follow-up email", "Reference one specific thing they said that connected to what you showed. Generic follow-ups get ignored."),
        ("Log their feedback", "What resonated, what was missing, any objections raised. This feeds directly into what we build next."),
        ("Answer any open questions directly", "If they asked about integrations, pricing, or data compliance — send a direct answer, not a deck."),
        ("Share login details", "Include the URL and demo login so they can explore independently: greg@workiflow.com / AscendaDemo!2026"),
    ]
    for action, detail in after_items:
        p2 = doc.add_paragraph()
        sp(p2, before=4, after=2)
        p2.paragraph_format.left_indent = Cm(0.3)
        r(p2, "▸  " + action + "  —  ", bold=True, size=10.5, color=VIOLET)
        r(p2, detail, size=10.5)

    doc.save("/Users/gregfranck/Ascenda/docs/guide-counsellor-presentation.docx")
    print("Saved counsellor guide")


# ═══════════════════════════════════════════════════════════════
# STUDENT GUIDE
# ═══════════════════════════════════════════════════════════════

def build_student():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # Title block
    p = doc.add_paragraph()
    sp(p, before=0, after=3)
    r(p, "Ascenda", bold=True, size=26, color=PURPLE)
    p = doc.add_paragraph()
    sp(p, before=0, after=2)
    r(p, "Student Presentation Guide", bold=True, size=13, color=VIOLET)
    p = doc.add_paragraph()
    sp(p, before=0, after=14)
    r(p, "1:1 buy-in  ·  ~15 min  ·  Feedback gathering", size=9.5, color=GREY)

    # ── PURPOSE ──
    section_heading(doc, "What this session is for", before=0)
    p = doc.add_paragraph()
    sp(p, before=6, after=4)
    r(p, "You're showing a student what the next year of their application process could look like — and then finding out where their real anxiety sits. Students don't want a product demo. They want to feel less overwhelmed.", size=10.5)
    p = doc.add_paragraph()
    sp(p, before=0, after=4)
    r(p, "Every feature you show should connect to a problem they've described. If you show something they didn't ask about, explain in one sentence why it matters for them specifically.", size=10.5)

    # ── BEFORE ──
    section_heading(doc, "Before the session")
    context_items = [
        ("What year are they in?", "Year 12 students care about building a list and understanding fit. Year 13 students care about deadlines, personal statement, and not dropping the ball."),
        ("IB or A-Level?", "Shapes which parts of the profile and matching feel relevant — IB students often have more subject flexibility concerns."),
        ("Do they have a shortlist already?", "If yes: the Timeline and match scores hit differently. If no: start with Explore and the Reach/Match/Safe framework."),
        ("Are they international?", "If yes: campus life section on the course page is critical. Lead with it explicitly."),
        ("Have they started their personal statement?", "If not: Essay Workshop is where the demo lands hardest. If yes: use it as a review tool, not a starting point."),
    ]
    p = doc.add_paragraph()
    sp(p, before=5, after=5)
    r(p, "Know these before you open the laptop:", size=10.5)
    for q, why in context_items:
        p2 = doc.add_paragraph()
        sp(p2, before=3, after=1)
        p2.paragraph_format.left_indent = Cm(0.3)
        r(p2, "▸  ", bold=True, size=10, color=VIOLET)
        r(p2, q + "  ", bold=True, size=10)
        r(p2, why, size=10, color=GREY)

    p = doc.add_paragraph()
    sp(p, before=8, after=3)
    r(p, "Technical setup:", bold=True, size=10)
    setup = [
        "Reset demo state so the platform starts clean — bell at zero.",
        "Platform open on /dashboard in a clean browser, zoomed in slightly (110%). Readability matters.",
        "If possible, have the demo profile's subject list reflect their subjects.",
    ]
    for s in setup:
        p2 = doc.add_paragraph()
        sp(p2, before=1, after=1)
        p2.paragraph_format.left_indent = Cm(0.3)
        r(p2, "▸  ", bold=True, size=10, color=VIOLET)
        r(p2, s, size=10)

    # ── OPENING ──
    section_heading(doc, "Opening — start with their stress (2 min)")
    p = doc.add_paragraph()
    sp(p, before=6, after=4)
    r(p, "Don't open with the platform. Open with a question that makes them feel heard:", size=10.5)

    pain_block(
        doc,
        pain_text="Students are overwhelmed by the application process but can't always articulate what's actually stressing them — they need to be asked.",
        action_text='Ask one: "What part of applying feels most unclear?" / "How confident do you feel about your university list right now?" / "What\'s taking up the most mental space?"',
        say_text='"I want to show you something a lot of students have found useful — not to do the work for you, but to make the whole thing feel less scattered. Takes about ten minutes. Tell me if anything doesn\'t make sense or seems off for your situation."',
    )
    note_line(doc, "Their answer tells you which sections to lean into. If they say 'I don't know what universities to apply to' → Explore. If they say 'I'm panicking about the personal statement' → Toolbox. If they say 'I have a list but I'm not sure it's right' → match scores + Timeline.")

    # ── PROFILE ──
    section_heading(doc, "Profile (1 min)")

    pain_block(
        doc,
        pain_text="Students don't realise their current subject choices might be quietly closing degree options — and by the time they do, it's often too late to change anything.",
        action_text="Click Profile (top nav). Show the filled-in profile. Point to the Pathway status pill.",
        say_text='"This is what a student fills in when they join — predicted grades, subjects, what they want to study. The reason we ask for all of this upfront: subject choices made years before you apply can quietly close doors without you realising. Roughly one in five students block themselves from degree options through decisions made earlier in school. This pill tells you whether your current path keeps your options open — and if something\'s wrong, it flags it while you can still do something about it."',
    )
    note_line(doc, "If they mention a specific subject they're uncertain about, reference it: 'So if you were to drop [subject], this is exactly where you'd see that flagged.'")

    # ── EXPLORE ──
    section_heading(doc, "Explore (2–3 min)")

    pain_block(
        doc,
        pain_text="Students don't know where they genuinely stand with universities — they either aim too safe or overreach, and they're guessing based on surface information.",
        action_text="Click Explore. Show Reach / Match / Safe cards with fit scores.",
        say_text='"This is what your shortlist looks like through Ascenda. The Reach / Match / Safe split — you\'ve probably heard that framework before. What makes this different is that it\'s not a guess: your profile gets scored against each university\'s actual entry requirements. And the score is transparent — you can see exactly why a course ranked where it did. If you disagree with something, you can see what it would take to change it."',
    )

    pain_block(
        doc,
        pain_text="International students are making life decisions about places they've never been — university websites are thin and inconsistent, and there's no consolidated source for what life is actually like.",
        action_text="Click into Cambridge Computer Science (Reach section). Scroll to the campus and city life section.",
        say_text='"When you click into a course, one page replaces thirty browser tabs. But the section I want to point out specifically — this one. If you\'re applying from outside the UK, it\'s genuinely hard to know what life is like at these places. International student percentage, cost of living, what the city is like, career prospects — all here. So when you\'re deciding between Manchester and Edinburgh, you\'re not making that call based on the university\'s own marketing."',
    )
    note_line(doc, "If they're a UK student, reframe: 'Even if you've been to these cities, this gives you a side-by-side comparison you can't easily get anywhere else.'")

    # ── TOOLBOX ──
    section_heading(doc, "Toolbox — essay workshop (2 min)")

    pain_block(
        doc,
        pain_text="Students have no benchmark for what a strong personal statement looks like — and the feedback they get from schools is inconsistent and often too late.",
        action_text='Click Toolbox → Essay Workshop. Clear the textarea. Type: "Science has always been my way of making sense of the world." Click Get Feedback.',
        say_text='"UCAS is 4,000 characters. Other portals are different lengths. And most students don\'t know what a strong personal statement actually looks like — the bar they\'re writing to is invisible. This benchmarks your draft against the strongest applications to Oxford, Cambridge, and Harvard. Not writing it for you — giving you immediate, structured feedback on what\'s working and what isn\'t, before it lands in front of a counsellor or an admissions office."',
    )
    note_line(doc, "If they haven't started: 'This is also useful before you write a word — the outline tool helps you build the structure from your profile.' If they've already written a draft: 'You can paste what you have in here right now.'")

    # ── APPLICATIONS ──
    section_heading(doc, "Applications (2 min)")

    pain_block(
        doc,
        pain_text="Managing multiple applications across different portals with different requirements and deadlines feels chaotic — students lose track of what's outstanding.",
        action_text="Click Applications. Show the What's next priority board at the top. Show one expanded application with its task list.",
        say_text='"This is the central hub. Everything in one place — not scattered across UCAS, university portals, and a spreadsheet. The priority board at the top shows the three most urgent things across all your applications: what to work on today, weighted by deadline pressure and how important the application is. Each application has its own task list so nothing falls through. You always know what\'s next."',
    )

    pain_block(
        doc,
        pain_text="When a student gets stuck on something specific, reaching their counsellor is clunky — emails get buried, context is lost by the time a meeting happens.",
        action_text="Click the violet Need help pill on the Cambridge card.",
        say_text='"If you get stuck — confused about requirements, not sure what to write, something\'s unclear — you click this. The platform already knows what you\'re working on and what stage you\'re at, so the request comes pre-filled with context. Your counsellor gets a structured message instead of a blank email. It\'s faster for you and more useful for them."',
    )

    # ── CLOSING ──
    section_heading(doc, "Closing — get their honest reaction (3 min)")
    p = doc.add_paragraph()
    sp(p, before=6, after=4)
    r(p, "Stop before they switch off. Don't end on a feature — end on a question:", size=10.5)

    tbl = doc.add_table(rows=1, cols=1)
    set_cell_bg(tbl.cell(0, 0), "F5F3FF")
    qp = tbl.cell(0, 0).paragraphs[0]
    qp.paragraph_format.space_before = Pt(7)
    qp.paragraph_format.space_after = Pt(7)
    qp.paragraph_format.left_indent = Cm(0.35)
    r(qp, "SAY:  ", bold=True, size=8.5, color=VIOLET)
    r(qp, '"That\'s the platform. I\'d rather hear your honest reaction than keep showing you things — what connected, and what didn\'t?"', italic=True, size=10.5, color=PURPLE)
    doc.add_paragraph()

    note_line(doc, "Don't defend features they push back on. That's the feedback. Note it and ask why.")

    # ── FEEDBACK QUESTIONS ──
    section_heading(doc, "Feedback questions to ask")
    p = doc.add_paragraph()
    sp(p, before=5, after=6)
    r(p, "Pick 2–3. Keep it conversational. The notes column is for you to fill in during the session.", size=10)

    two_col_table(doc, [
        ("Is there a part of the application process you'd actually use this for?", ""),
        ("What's missing for how you personally work?", ""),
        ("Does the match score feel useful, or does it feel like a black box?", ""),
        ("Would you use this alongside your counsellor or more on your own?", ""),
        ("Is there anything here that would make you feel less stressed?", ""),
        ("What would you show a friend if you were recommending this?", ""),
        ("Is the personal statement feedback the kind of thing you'd act on?", ""),
        ("What would need to be different for you to use this every week?", ""),
    ], "Question to ask", "Notes")

    # ── WHAT TO LISTEN FOR ──
    section_heading(doc, "What to listen for")
    p = doc.add_paragraph()
    sp(p, before=5, after=6)
    r(p, "These signals tell you whether the buy-in worked — and where the real friction is.", size=10)

    tbl = doc.add_table(rows=5, cols=2)
    tbl.style = 'Table Grid'
    set_cell_bg(tbl.cell(0, 0), "ECFDF5")
    set_cell_bg(tbl.cell(0, 1), "FEF2F2")
    tbl.columns[0].width = Cm(8.2)
    tbl.columns[1].width = Cm(7.6)
    r(tbl.cell(0, 0).paragraphs[0], "They're bought in if they say…", bold=True, size=9, color=GREEN)
    r(tbl.cell(0, 1).paragraphs[0], "Worth digging into if they say…", bold=True, size=9, color=RED)
    signals = [
        ('"This would have helped me with…"', '"I don\'t really understand how the score works"'),
        ('"Can I actually sign up for this?"', '"My counsellor doesn\'t use anything like this"'),
        ('"I would have used this when I was doing [X]"', '"I already have a system that works for me"'),
        ('"The personal statement feedback looks really useful"', '"This feels like a lot to keep updated"'),
    ]
    for i, (pos, neg) in enumerate(signals, 1):
        tbl.cell(i, 0).width = Cm(8.2)
        tbl.cell(i, 1).width = Cm(7.6)
        r(tbl.cell(i, 0).paragraphs[0], pos, italic=True, size=9.5)
        r(tbl.cell(i, 1).paragraphs[0], neg, italic=True, size=9.5)

    p = doc.add_paragraph()
    sp(p, before=0, after=10)

    # ── AFTER ──
    section_heading(doc, "After the session — within 24 hours")
    after_items = [
        ("Share the platform URL and login", "Include: ascenda-ashy.vercel.app / greg@workiflow.com / AscendaDemo!2026. Students who can explore independently will."),
        ("Note what resonated", "Specifically: which part of the application process they said felt most overwhelming. This is the hook for the follow-up."),
        ("Answer open questions directly", "If they asked something you couldn't answer on the spot, follow up with the specific answer. Students remember this."),
        ("Reference something they said", "A follow-up that quotes their own words — 'You mentioned you were worried about X' — lands much better than a generic message."),
    ]
    for action, detail in after_items:
        p2 = doc.add_paragraph()
        sp(p2, before=4, after=2)
        p2.paragraph_format.left_indent = Cm(0.3)
        r(p2, "▸  " + action + "  —  ", bold=True, size=10.5, color=VIOLET)
        r(p2, detail, size=10.5)

    doc.save("/Users/gregfranck/Ascenda/docs/guide-student-presentation.docx")
    print("Saved student guide")


build_counsellor()
build_student()
print("Done.")
